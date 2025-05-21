import streamlit as st
import datetime
import fitz  # PyMuPDF
import json
import paramiko
import os
import getpass
from transformers import pipeline, logging as tf_logging
from fpdf import FPDF
import spacy
import warnings
import tensorflow as tf
import pandas as pd
import logging
import tempfile
from groq import Groq
import joblib
from sklearn.pipeline import Pipeline
from sklearn.naive_bayes import MultinomialNB
from sklearn.feature_extraction.text import TfidfVectorizer
from docx import Document  # Add this import for Word document support

# Suppress warnings
warnings.filterwarnings('ignore')
tf_logging.set_verbosity(tf_logging.ERROR)
tf.compat.v1.logging.set_verbosity(tf.compat.v1.logging.ERROR)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("egov_grievance.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Initialize Groq client with error handling
try:
    groq_client = Groq(api_key="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")
except Exception as e:
    logger.error(f"Failed to initialize Groq client: {str(e)}")
    groq_client = None

# Mainframe connection settings
HOSTNAME = "204.90.115.200"
LOCAL_FILE = "classified_complaints.txt"
REMOTE_FILE = "/z/z67427/classified_complaints.txt"

# Model and category-priority map (from classify1.py)
def train_model():
    data = pd.read_csv("sample_complaints_100.csv")
    X = data["Complaint Text"]
    y = data["Category"]
    pipeline = Pipeline([
        ('tfidf', TfidfVectorizer()),
        ('nb', MultinomialNB())
    ])
    pipeline.fit(X, y)
    joblib.dump(pipeline, "complaint_classifier.pkl")
    return pipeline
try:
    model = joblib.load("complaint_classifier.pkl")
except:
    model = train_model()
category_priority_map = {
    "Bank": "06",
    "Electricity": "05",
    "Roads": "04",
    "Water": "03",
    "Sanitation": "02",
    "Other": "01"
}

def init_session_state():
    if 'page' not in st.session_state:
        st.session_state.page = 'login'
    if 'username' not in st.session_state:
        st.session_state.username = ''
    if 'password' not in st.session_state:
        st.session_state.password = ''
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'login_time' not in st.session_state:
        st.session_state.login_time = None
    if 'session_expiry' not in st.session_state:
        st.session_state.session_expiry = None
    if 'classification_done' not in st.session_state:
        st.session_state.classification_done = False
    if 'nlp_model' not in st.session_state:
        st.session_state.nlp_model = None
    if 'groq_initialized' not in st.session_state:
        st.session_state.groq_initialized = False
        try:
            # Test Groq connection
            groq_client.chat.completions.create(
                messages=[{"role": "user", "content": "test"}],
                model="llama-3.3-70b-versatile",
                max_tokens=10
            )
            st.session_state.groq_initialized = True
        except Exception as e:
            logger.error(f"Failed to initialize Groq: {str(e)}")
            st.session_state.groq_initialized = False
    if 'processed_df' not in st.session_state:
        st.session_state.processed_df = None
    if 'pdf_data' not in st.session_state:
        st.session_state.pdf_data = None
    if 'show_download' not in st.session_state:
        st.session_state.show_download = False

def check_session_validity():
    if not st.session_state.authenticated:
        return False
    
    if not st.session_state.login_time or not st.session_state.session_expiry:
        return False
    
    current_time = datetime.datetime.now()
    if current_time > st.session_state.session_expiry:
        # Session expired, clear session
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        init_session_state()
        return False
    
    return True

@st.cache_resource
def load_nlp_model():
    return spacy.load('en_core_web_sm')

def verify_credentials(username, password):
    try:
        ssh = paramiko.SSHClient()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        ssh.connect(hostname=HOSTNAME, username=username, password=password)
        ssh.close()
        return True
    except:
        return False

def login_page():
    st.title("🔐 Mainframe Authentication")
    
    # Check if there's a valid session
    if check_session_validity():
        st.session_state.page = 'classifier'
        st.rerun()
        return
    
    col1, col2 = st.columns([3, 2])
    
    with col1:
        username = st.text_input(
            "Username (z-ID)", 
            placeholder="z#####",
            value=st.session_state.username
        )
        
        password = st.text_input(
            "Password", 
            type="password",
            placeholder="######## (8-digit mainframe password)",
            value=st.session_state.password
        )

    with col2:
        st.markdown("""
        ### Instructions
        1. Enter your z-ID (starts with 'z')
        2. Enter your 8-digit mainframe password
        3. Click Login to proceed
        """)

    if st.button("Login"):
        if not username or not password:
            st.error("Please enter both username and password!")
            return
        
        if not username.startswith('z') or len(username) < 6:
            st.error("Invalid z-ID format! Should be 'z' followed by 5 digits.")
            return
            
        if len(password) < 8:
            st.error("Invalid password format! Should be 8 digits.")
            return

        with st.spinner("Verifying credentials..."):
            if verify_credentials(username, password):
                st.session_state.username = username
                st.session_state.password = password
                st.session_state.authenticated = True
                st.session_state.login_time = datetime.datetime.now()
                # Set session expiry to 8 hours from login
                st.session_state.session_expiry = st.session_state.login_time + datetime.timedelta(hours=8)
                st.session_state.page = 'classifier'
                st.success("Login successful!")
                st.rerun()
            else:
                st.error("Invalid credentials!")

# Status generator
def generate_status(date_str):
    try:
        complaint_date = datetime.datetime.strptime(date_str, "%Y%m%d")
        days_old = (datetime.datetime.now() - complaint_date).days
        return "E" if days_old > 5 else "O"
    except Exception as e:
        st.error(f"Error generating status: {e}")
        return "Unknown"

# Extract text from uploaded PDF
def extract_text_from_pdf(uploaded_file):
    try:
        text = ""
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text("text")
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

# Extract text from uploaded Word document
def extract_text_from_word(uploaded_file):
    try:
        doc = Document(uploaded_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from Word document: {e}")
        return ""

def upload_to_mainframe():
    try:
        st.info("🚀 Connecting to mainframe...")
        remote_file = f"/z/{st.session_state.username}/classified_complaints.txt"

        ssh = paramiko.SSHClient()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        ssh.connect(
            hostname=HOSTNAME, 
            username=st.session_state.username, 
            password=st.session_state.password
        )

        st.success("🔗 Connected successfully.")

        sftp = ssh.open_sftp()
        st.info(f"📤 Uploading {LOCAL_FILE} to {remote_file} (after converting to EBCDIC)...")

        with open(LOCAL_FILE, 'r', encoding='utf-8') as f:
            utf8_text = f.read()

        utf8_text = utf8_text.replace('\n', '\r')
        ebcdic_data = utf8_text.encode('cp037')

        with sftp.file(remote_file, mode='wb') as remote_file_obj:
            remote_file_obj.write(ebcdic_data)

        sftp.close()
        ssh.close()

        st.success(f"✅ Upload successful! File has been converted to EBCDIC and uploaded to mainframe")
        return True

    except Exception as e:
        st.error(f"❌ Upload failed: {str(e)}")
        return False

# Word document based classification
def word_based_classify(complaint):
    try:
        doc = Document(complaint)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text.lower() + " "
        
        if "power" in text or "electric" in text or "light" in text:
            return "Electricity", "05"
        elif "pothole" in text or "road" in text or "traffic" in text:
            return "Roads", "04"
        elif "water" in text or "supply" in text:
            return "Water", "03"
        elif "garbage" in text or "sanitation" in text:
            return "Sanitation", "02"
        elif "bank" in text:
            return "Bank", "06"
        else:
            return "Other", "01"
    except Exception as e:
        st.error(f"Error processing Word document: {e}")
        return "Other", "01"

def classifier_page():
    st.title("📄 E-Governance Complaint Classifier")
    st.write("Upload a PDF or Word document containing complaints (one per line or separated clearly).")
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])
    if uploaded_file is not None:
        # Extract text based on file type
        if uploaded_file.name.endswith('.pdf'):
            pdf_text = extract_text_from_pdf(uploaded_file)
            if not pdf_text:
                st.error("No text extracted from the PDF.")
                return
            complaints = [line.strip() for line in pdf_text.split("\n") if line.strip()]
        else:  # .docx file
            word_text = extract_text_from_word(uploaded_file)
            if not word_text:
                st.error("No text extracted from the Word document.")
                return
            complaints = [line.strip() for line in word_text.split("\n") if line.strip()]
        base_dates = [
            (datetime.datetime.now() - datetime.timedelta(days=i)).strftime("%Y%m%d")
            for i in range(len(complaints))
        ]
        output_data = []
        for idx, (complaint, date) in enumerate(zip(complaints, base_dates), start=1):
            comp_id = f"GRV{idx:07d}"
            # Try model-based classification, fallback to word-based if error
            try:
                category = model.predict([complaint])[0]
                priority = category_priority_map.get(category, "01")
            except Exception:
                category, priority = word_based_classify(complaint)
            status = generate_status(date)
            complaint_entry = {
                "Complaint_ID": comp_id,
                "Complaint_Text": complaint,
                "Priority": priority,
                "Date": date,
                "Status": status,
                "Category": category
            }
            output_data.append(complaint_entry)
        text_data = "\n".join(
            [
                f"{complaint['Complaint_ID']} {complaint['Category']}".ljust(23) 
                + f"{complaint['Complaint_Text']}".ljust(57)
                + f"{complaint['Priority']} {complaint['Date']} {complaint['Status']}"
                for complaint in output_data
            ]
        )
        st.success(f"Processed {len(complaints)} complaints!")
        with open(LOCAL_FILE, 'w', encoding='utf-8') as f:
            f.write(text_data)
        st.subheader("Preview of Classified Complaints")
        for complaint in output_data:
            preview_line = (
                f"{complaint['Complaint_ID']} {complaint['Category']}".ljust(23) 
                + f"{complaint['Complaint_Text']}".ljust(57)
                + f"{complaint['Priority']} {complaint['Date']} {complaint['Status']}"
            )
            st.text(preview_line)
        col1, col2 = st.columns([1, 2])
        with col1:
            if st.button("Upload to Mainframe"):
                with st.spinner("Uploading to mainframe..."):
                    if upload_to_mainframe():
                        st.session_state.classification_done = True
                        st.balloons()
                        try:
                            os.remove(LOCAL_FILE)
                            st.info("Local temporary file cleaned up.")
                        except:
                            st.warning("Could not clean up local temporary file.")
                        st.success("Ready to proceed to Grievance Analysis!")
        with col2:
            if st.session_state.classification_done:
                if st.button("Proceed to Grievance Analysis"):
                    st.session_state.page = 'analysis'
                    st.rerun()

def fetch_grievance_data():
    try:
        ssh = paramiko.SSHClient()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        ssh.connect(
            hostname=HOSTNAME, 
            username=st.session_state.username, 
            password=st.session_state.password
        )
        sftp = ssh.open_sftp()
        remote_file = f"/z/{st.session_state.username}/grievance_data.json"
        local_file = 'grievance_data.json'
        sftp.get(remote_file, local_file)
        sftp.close()
        ssh.close()
        logger.info("Downloaded grievance_data.json from USS")
        
        with open(local_file, 'r') as f:
            data = json.load(f)
        return data
    except Exception as e:
        logger.exception("Failed to fetch grievance data")
        st.error(f"Failed to fetch grievance data: {str(e)}")
        return []

def clean_text(text, nlp):
    doc = nlp(text)
    return " ".join([token.text for token in doc if not token.is_punct])

def process_grievances(texts, nlp, batch_size=10):
    solutions = []
    total = len(texts)
    progress_text = "Generating solutions..."
    
    # Create a cache dictionary to store solutions
    solution_cache = {}
    
    my_bar = st.progress(0, text=progress_text)
    
    # Process in larger batches
    for i in range(0, total, batch_size):
        batch = texts[i:i + batch_size]
        batch_solutions = []
        
        # Check cache first
        for text in batch:
            if text in solution_cache:
                batch_solutions.append(solution_cache[text])
            else:
                solution = generate_solution(text, nlp)
                solution_cache[text] = solution
                batch_solutions.append(solution)
            
            current_count = len(solutions) + len(batch_solutions)
            progress = min(1.0, current_count / total)
            my_bar.progress(progress, text=f"{progress_text} ({current_count}/{total})")
        
        solutions.extend(batch_solutions)
        
    my_bar.empty()
    return solutions

# Generate AI solution with improved prompt
def generate_solution(text, _nlp):
    if not st.session_state.groq_initialized:
        return "Error: Groq API not initialized"
    
    if not text.strip():
        return "No grievance description provided."
    
    try:
        cleaned = clean_text(text, _nlp)
        # Enhanced prompt for structured solutions
        prompt = (
            f"Analyze this complaint and provide a structured solution: '{cleaned}'\n\n"
            "Format your response in the following structure:\n"
            "1. Issue Summary: Brief description of the problem\n"
            "2. Immediate Actions: List of immediate steps to take\n"
            "3. Long-term Solutions: Sustainable measures to prevent recurrence\n"
            "4. Department Requirements: Specific resources or actions needed from departments\n"
            "5. Safety Measures: Required safety precautions\n"
            "Keep each section concise and actionable. Use bullet points for clarity."
        )
        
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "You are a government grievance resolution specialist. Provide clear, structured solutions that can be easily implemented by different departments. Focus on practical, actionable steps and maintain a professional tone."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            model="llama-3.3-70b-versatile",
            temperature=0.7,
            max_tokens=400,  # Increased slightly for better structure
            top_p=0.9,
            stream=False
        )
        
        solution = chat_completion.choices[0].message.content.strip()
        
        # Format the solution for better display
        formatted_solution = solution.replace('\n\n', '\n').replace('•', '• ').replace('*', '')
        # Ensure proper spacing after bullet points
        formatted_solution = '\n'.join(
            line if line.strip().startswith('•') else f"  {line}"
            for line in formatted_solution.split('\n')
        )
        
        return formatted_solution
        
    except Exception as e:
        logger.exception(f"Error generating solution: {str(e)}")
        if not st.session_state.groq_initialized:
            st.session_state.groq_initialized = False  # Mark as failed for retry
        return "Error generating solution. Please try again later."

class PDFReport(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        self.set_margins(10, 10, 10)
        
    def header(self):
        self.set_font("Arial", "B", 14)
        self.cell(0, 10, "E-Governance Grievance Report (AI Solutions)", ln=True, align="C")
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

    def add_grievance(self, row):
        self.set_font("Arial", "B", 10)
        self.cell(0, 5, f"Grievance ID: {row.ID}", ln=True)
        self.set_font("Arial", "", 10)
        self.cell(0, 5, f"Priority: {row.Priority} (Adjusted: {row['Adjusted Priority']}) | Age: {row['Age (days)']} days | Status: {row.Status}", ln=True)
        
        self.set_font("Arial", "B", 10)
        self.cell(0, 5, "Issue:", ln=True)
        self.set_font("Arial", "", 10)
        self.multi_cell(0, 5, str(row.Text))
        
        self.set_font("Arial", "B", 10)
        self.cell(0, 5, "Solution:", ln=True)
        self.set_font("Arial", "", 10)
        self.multi_cell(0, 5, str(row.Solution))
        self.ln(5)

def generate_and_download_pdf():
    if st.session_state.processed_df is None:
        st.error("No data available for PDF generation!")
        return
    
    try:
        df_sorted = st.session_state.processed_df
        categories = df_sorted['Category'].unique().tolist()
        
        pdf = PDFReport()
        for cat in categories:
            sub = df_sorted[df_sorted['Category'] == cat]
            if sub.empty:
                continue
            pdf.add_page()
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, f"Category: {cat}", ln=True)
            pdf.ln(2)
            for _, row in sub.iterrows():
                pdf.add_grievance(row)
        
        path = os.path.join(tempfile.gettempdir(), "grievance_report.pdf")
        pdf.output(path)
        
        with open(path, 'rb') as f:
            pdf_data = f.read()
        
        os.remove(path)  # Clean up temp file
        return pdf_data
    except Exception as e:
        logger.exception("Error generating PDF")
        st.error(f"Failed to generate PDF report: {str(e)}")
        return None

def analysis_page():
    st.title("📊 E-Governance Grievance Analyzer")
    
    # Check Groq initialization
    if not st.session_state.groq_initialized:
        st.error("⚠️ Groq API is not properly initialized. Please check your API key and try again.")
        if st.button("Retry Groq Connection"):
            try:
                groq_client.chat.completions.create(
                    messages=[{"role": "user", "content": "test"}],
                    model="llama-3.3-70b-versatile",
                    max_tokens=10
                )
                st.session_state.groq_initialized = True
                st.success("✅ Groq connection restored!")
                st.rerun()
            except Exception as e:
                st.error(f"Failed to connect to Groq: {str(e)}")
        return

    if 'nlp_model' not in st.session_state or st.session_state.nlp_model is None:
        with st.spinner("Loading NLP model..."):
            st.session_state.nlp_model = load_nlp_model()

    # Only fetch and process data if not already in session state
    if st.session_state.processed_df is None:
        data = fetch_grievance_data()
        if not data:
            st.warning("No grievance data available.")
            return

        df = pd.DataFrame(data)
        df = df.fillna({
            "Category": "Unknown",
            "Text": "",
            "Priority": 0,
            "Status": "Unknown",
            "Date": "",
            "Age (days)": 0
        })

        if 'Solution' not in df.columns or df['Solution'].isnull().all():
            with st.spinner("Generating AI Solutions..."):
                texts = df['Text'].tolist()
                solutions = process_grievances(texts, st.session_state.nlp_model)
                df['Solution'] = solutions

        df['Adjusted Priority'] = df.apply(
            lambda row: min(int(row['Priority'] or 0) + (3 if row.get('Age (days)', 0) > 90 else 2 if row.get('Age (days)', 0) > 60 else 1 if row.get('Age (days)', 0) > 30 else 0), 10),
            axis=1
        )
        
        st.session_state.processed_df = df.sort_values(by=['Adjusted Priority', 'Age (days)'], ascending=[False, False])

    # Display the processed dataframe
    df_sorted = st.session_state.processed_df
    st.subheader("📋 Prioritized Grievance List by Category")
    categories = df_sorted['Category'].unique().tolist()
    
    for cat in categories:
        sub = df_sorted[df_sorted['Category'] == cat]
        if sub.empty:
            continue
            
        st.markdown(f"### Category: **{cat}**")
        
        display_df = sub[["ID", "Priority", "Adjusted Priority", "Age (days)", "Status", "Date", "Text", "Solution"]].copy()
        
        # Configure the Solution column to display formatted text
        st.dataframe(
            display_df,
            column_config={
                "Solution": st.column_config.TextColumn(
                    "Solution",
                    width="large",
                    help="AI-generated solution for the grievance"
                ),
                "Text": st.column_config.TextColumn(
                    "Complaint",
                    width="medium",
                    help="Original complaint text"
                )
            },
            hide_index=True
        )

    # PDF Generation and Download
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("📄 Generate PDF Report", type="primary"):
            with st.spinner("Preparing PDF report..."):
                pdf_data = generate_and_download_pdf()
                if pdf_data:
                    st.session_state.pdf_data = pdf_data
                    st.session_state.show_download = True
                    st.success("PDF Generated Successfully! 🎉")
                    st.balloons()
    
    with col2:
        if st.session_state.get('show_download', False):
            if st.download_button(
                "⬇️ Download PDF Report",
                data=st.session_state.pdf_data,
                file_name="Grievance_Report.pdf",
                mime="application/pdf",
            ):
                st.success("Download completed! 🌟")
                st.snow()  # Add a snow animation on successful download

def main():
    st.set_page_config(layout="wide")
    init_session_state()
    
    # Check session validity at the start
    if st.session_state.authenticated and not check_session_validity():
        st.warning("Your session has expired. Please log in again.")
        st.session_state.page = 'login'
    
    # Sidebar for navigation when authenticated
    if st.session_state.authenticated:
        with st.sidebar:
            st.success(f"Logged in as: {st.session_state.username}")
            # Show session expiry time
            if st.session_state.session_expiry:
                remaining_time = st.session_state.session_expiry - datetime.datetime.now()
                hours, remainder = divmod(remaining_time.seconds, 3600)
                minutes, _ = divmod(remainder, 60)
                st.info(f"Session expires in: {hours}h {minutes}m")
            
            if st.button("Logout"):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                init_session_state()
                st.rerun()
    
    # Page routing
    if st.session_state.page == 'login':
        login_page()
    elif st.session_state.page == 'classifier':
        classifier_page()
    elif st.session_state.page == 'analysis':
        analysis_page()

if __name__ == "__main__":
    main() 
